from docx import Document
from io import BytesIO


def create_resume_docx(content):

    doc = Document()

    doc.add_heading(
        "Tailored Resume",
        level=1
    )

    doc.add_paragraph(content)

    file_stream = BytesIO()

    doc.save(file_stream)

    file_stream.seek(0)

    return file_stream
